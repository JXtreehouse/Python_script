import os
# Python docx module allows user to manipulate docs by either manipulating the existing one or creating a new empty
# document and manipulating it. It is a powerful tool as it helps you to manipulate the document to a very large
# extend.
from docx import Document
from docxcompose.composer import Composer


def auto_merge_doc(source_file_path_list, target_file_path):
    # 1. Create docx object
    page_break_doc = Document()
    # To add a page break in a word document you can use add_page_break() method. This method adds a hard page break
    # in the document and creates a new paragraph object. This is an essential paragraph styling tool. Whenever the
    # requirement is to start a new page, mostly cases where you want a new page for a new topic, a single method is
    # sufficient. It helps increase clarity and improves presentation to reap the most of word.
    page_break_doc.add_page_break()

    target_doc = Document(source_file_path_list[0])
    target_composer = Composer(target_doc)
    for index in range(len(source_file_path_list)):
        # 跳过第一个作为模板的文件
         if index ==0:
             continue
         # 填充分页符文档
         target_composer.append(page_break_doc)
         # 拼接文档内容
         f = source_file_path_list[index]
         target_composer.append(Document(f))
         # 保存目标文档
         target_composer.save(target_file_path)

if __name__ == '__main__':
    source_path = os.getcwd() + '/data/'
    target_file = os.getcwd() + '/data/merge_result.docx'
    source_file_list = os.listdir(source_path)
    source_file_list_all = []
    for file in source_file_list:
        source_file_list_all.append(source_path + file)
    auto_merge_doc(source_file_list_all, target_file)
