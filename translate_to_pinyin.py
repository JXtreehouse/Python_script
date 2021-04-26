'''
Author: AlexZ33
Date: 2021-04-26 09:42:02
LastEditTime: 2021-04-26 19:33:23
LastEditors: Please set LastEditors
Description: In User Settings Edit
FilePath: /Python_script/translate_to_pinyin.py
'''
# 读取docx中的文本并转化成拼音
import os
import re
import docx
import pypinyin

path = os.getcwd()

inputPath = path + '/data/飞花令整理版.docx'
outputfile = path + '/data/tmp/飞花令拼音版.docx'

if __name__ == "__main__":
    # 获取文档对象
    file = docx.Document(inputPath)
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为448，每个回车隔离一段
    # 输出每一段的内容
    document = docx.Document()
    for p in file.paragraphs:
        txt = p.text


        # pinyin 方法默认使用了 TONE 的风格，而 lazy_pinyin 方法默认使用了 NORMAL 的风格，所以就导致二者返回风格不同了
        str = pypinyin.lazy_pinyin(txt, style='Style.TONE')
        split_str = str[1: -1]
        s = ''
        for i in split_str:

            s = s + ''.join(i) + " "

        print(txt)
        print(s)
        # document = docx.Document()

        document.add_paragraph(s)
        document.add_paragraph(txt)
    document.save(outputfile)
